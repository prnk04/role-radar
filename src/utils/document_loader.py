"""
Extract text from various documents
"""

from pathlib import Path
from typing import Optional

# import PyPDF2
import pypdf
import docx
import io
import requests


class DocumentLoader:
    """Load and extract text from documents"""

    @staticmethod
    def load_pdf(file_path: Path):
        """Extract text from pdf"""
        print("should extract text from pdf")

        try:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)

                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                print("ext: ", text)
                return text.strip()

        except Exception as e:
            print("Exception in loading pdf: ", e)
            raise Exception(f"Error loading PDF: {e}")

    @staticmethod
    def load_docx(file_path: str):
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
        except Exception as e:
            raise Exception(f"Error loading DOCX: {e}")

    @staticmethod
    def load_txt(file_path: Path) -> str:
        """Load text file"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"Error loading TXT: {e}")

    @classmethod
    def load_from_url(cls, url: str):
        """Load file from url"""
        try:
            url_res = requests.get(url)

            content_type = url_res.headers["content-type"]
            match content_type.split(";")[0]:
                case "application/pdf":
                    reader = pypdf.PdfReader(io.BytesIO(url_res.content))
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return {
                        "file_type": "PDF",
                        "size": 0,
                        "text": text.strip(),
                        "word_count": len(text.split(" ")),
                    }

                case "text/html":
                    # return url_res.text.strip()
                    text = url_res.text.strip()
                    if text.startswith("<!doctype html"):
                        raise Exception("Invalid file format")
                    return {
                        "file_type": "TXT",
                        "size": 0,
                        "text": url_res.text.strip(),
                        "word_count": len(url_res.text.split(" ")),
                    }
                case "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    # load docx file
                    file_stream = io.BytesIO(url_res.content)
                    doc = docx.Document(file_stream)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return {
                        "file_type": "DOCX",
                        "size": 0,
                        "text": text.strip(),
                        "word_count": len(text.split(" ")),
                    }
                case _:
                    print(
                        "Unknown format. Please use URLs that have pdf/text or docx files only"
                    )
                    # return {
                    #     "file_type": "Unknown",
                    #     "size": 0,
                    #     "text": "Unknown format. Please use URLs that have pdf/text or docx files only",
                    #     "word_count": 0,
                    # }
                    raise Exception("Invalid file format")
        except Exception as e:
            raise Exception(f"InvalidFileFormat: {e}")

    @classmethod
    def load(cls, file_path: Path) -> dict:
        """
        Load document and return metadata + text

        Returns:
            dict with 'text', 'filename', 'file_type', 'size'
        """
        file_path = Path(file_path)
        print("file path: ", file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        print("file exists")

        # Determine file type
        suffix = file_path.suffix.lower()

        print("suffix:L ", suffix == ".pdf")
        # Extract text based on type
        if suffix == ".pdf":
            print("I should go")
            text = cls.load_pdf(file_path)
            file_type = "PDF"
        elif suffix in [".docx", ".doc"]:
            text = cls.load_docx(str(file_path))
            file_type = "DOCX"
        elif suffix == ".txt":
            text = cls.load_txt(file_path)
            file_type = "TXT"
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        return {
            "text": text,
            "filename": file_path.name,
            "file_type": file_type,
            "size": file_path.stat().st_size,
            "word_count": len(text.split()),
        }
